from docx import Document

# Create the document

doc = Document()
doc.add_heading('SkyrakSys HRM Production Deployment Fix Script', 0)

sections = [
    ('Introduction & Purpose', [
        ('Script Name', 'complete-deployment-fix.sh'),
        ('Purpose', 'Automated fix for deployment issues in SkyrakSys HRM production'),
        ('Issues Addressed', 'DB migration tracking, missing indexes, PM2 config, service health, security'),
        ('Usage', 'Run on production server after pulling latest code')
    ]),
    ('Environment Preparation', [
        ('Tool Checks', 'Verifies Node.js, PostgreSQL, PM2 installed'),
        ('Env Validation', 'Checks .env and .env.production for required variables'),
        ('Permissions', 'Ensures script has sudo/root access')
    ]),
    ('Database Validation & Fixes', [
        ('DB Connection', 'Validates DB connection and credentials'),
        ('Table Checks', 'Checks for missing tables (SequelizeMeta, etc.)'),
        ('Indexes', 'Adds missing indexes for performance')
    ]),
    ('Migration Tracking', [
        ('Migration Status', 'Ensures migration history is correct'),
        ('Meta Table', 'Populates SequelizeMeta if missing'),
        ('Verification', 'Verifies migration status and logs')
    ]),
    ('Application Services', [
        ('PM2 Restart', 'Restarts backend and frontend services via PM2'),
        ('Config Update', 'Updates PM2 configuration if needed'),
        ('Health Check', 'Checks service health and logs errors')
    ]),
    ('Frontend & Backend Health Checks', [
        ('API Endpoints', 'Verifies API endpoints are reachable'),
        ('Frontend Status', 'Checks frontend availability'),
        ('Error Logging', 'Logs errors and warnings for review')
    ]),
    ('Security & Permissions', [
        ('File Permissions', 'Ensures file and directory permissions'),
        ('Header Checks', 'Validates CORS, COOP, and other headers'),
        ('Debug Panel', 'Checks for exposed debug panels')
    ]),
    ('Final Validation & Reporting', [
        ('Summary', 'Summarizes actions taken'),
        ('Troubleshooting', 'Provides troubleshooting tips'),
        ('Next Steps', 'Lists verification commands and recommendations')
    ])
]

for i, (title, items) in enumerate(sections, 1):
    doc.add_page_break()
    doc.add_heading(f'Page {i}: {title}', level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Step/Section'
    hdr_cells[1].text = 'Details'
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]

# Save the document

doc.save(r'd:\skyraksys_hrm\redhatprod\scripts\complete-deployment-fix-summary.docx')
